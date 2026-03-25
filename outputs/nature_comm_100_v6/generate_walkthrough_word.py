"""Generate CrossDisc_Bench pipeline walkthrough Word document using v6 data."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, os

BASE = os.path.dirname(os.path.abspath(__file__))

# ── Load v6 data ──
with open(os.path.join(BASE, "extraction_results.json"), encoding="utf-8") as f:
    extraction = json.load(f)
with open(os.path.join(BASE, "benchmark_dataset.json"), encoding="utf-8") as f:
    benchmark = json.load(f)
with open(os.path.join(BASE, "p1p5_comparison_summary.json"), encoding="utf-8") as f:
    p1p5 = json.load(f)
with open(os.path.join(BASE, "p5_kg_eval_results.json"), encoding="utf-8") as f:
    kg_eval = json.load(f)

# Get first paper data
paper0 = extraction[0]
gt0 = benchmark[0]
kg0 = kg_eval[0] if kg_eval else None

doc = Document()

# ── Global styles ──
style = doc.styles["Normal"]
style.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5fae\u8f6f\u96c5\u9ed1")
    hs.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    hs.paragraph_format.space_after = Pt(6)

# ── Helper functions ──

def add_code_block(text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        run.element.rPr.append(shd)

def add_json_block(obj, indent=2):
    if isinstance(obj, str):
        add_code_block(obj)
    else:
        add_code_block(json.dumps(obj, ensure_ascii=False, indent=indent))

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)

def add_bold_text(bold_part, normal_part):
    p = doc.add_paragraph()
    run_b = p.add_run(bold_part)
    run_b.bold = True
    p.add_run(normal_part)

def add_card_box(title, content, color=RGBColor(0x1A, 0x56, 0xDB)):
    """Add a colored card-style box with title and content."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    # Title with color background
    run_t = p.add_run(f"  {title}  ")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    run_t.element.rPr.append(shd)
    # Content
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.add_run(content).font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════
#  DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════

# ── Title ──
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(4)
run = title_p.add_run("CrossDisc-Bench Pipeline Walkthrough")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("\u8de8\u5b66\u79d1\u5047\u8bbe\u751f\u6210\u57fa\u51c6\u6d4b\u8bd5 \u2014 \u5168\u6d41\u7a0b\u5b9e\u4f8b\u5c55\u793a (v6)")
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run(f"\u6570\u636e\u6765\u6e90: Nature Communications | \u5904\u7406\u8bba\u6587: {len(extraction)} \u7bc7 | GT\u6761\u76ee: {len(benchmark)} \u6761").font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# Section 1: Pipeline Overview
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Pipeline Overview (\u6d41\u7a0b\u603b\u89c8)", level=1)

doc.add_paragraph(
    "CrossDisc-Bench \u662f\u4e00\u4e2a\u9762\u5411\u8de8\u5b66\u79d1\u5047\u8bbe\u751f\u6210\u7684\u57fa\u51c6\u6d4b\u8bd5\u7cfb\u7edf\u3002"
    "\u5b83\u4ece Nature Communications \u8bba\u6587\u51fa\u53d1\uff0c\u901a\u8fc7\u591a\u9636\u6bb5 LLM \u7ba1\u7ebf\u6784\u5efa\u77e5\u8bc6\u56fe\u8c31\u3001"
    "\u751f\u6210\u8de8\u5b66\u79d1\u5047\u8bbe\u8def\u5f84\uff0c\u5e76\u6784\u5efa\u57fa\u4e8e\u8bc1\u636e\u7684 Ground Truth \u6570\u636e\u96c6\u3002"
)

add_table(
    ["\u9636\u6bb5", "\u540d\u79f0", "\u8f93\u5165", "\u8f93\u51fa"],
    [
        ["Stage 1", "\u5b66\u79d1\u5206\u7c7b", "\u8bba\u6587\u6807\u9898+\u6458\u8981", "\u4e3b\u5b66\u79d1+\u526f\u5b66\u79d1"],
        ["Stage 2a", "\u6982\u5ff5\u62bd\u53d6", "\u8bba\u6587+\u5b66\u79d1\u5206\u7c7b", "\u591a\u5b66\u79d1\u6982\u5ff5\u5217\u8868"],
        ["Stage 2b", "\u5173\u7cfb\u62bd\u53d6", "\u6982\u5ff5\u5217\u8868", "\u5b66\u79d1\u95f4\u5173\u7cfb\u4e09\u5143\u7ec4"],
        ["Stage 2c", "\u67e5\u8be2\u751f\u6210", "\u6982\u5ff5+\u5173\u7cfb", "L1/L2/L3 \u5c42\u6b21\u67e5\u8be2"],
        ["Stage 2d", "\u5047\u8bbe\u8def\u5f84\u751f\u6210", "\u67e5\u8be2+KG", "L1/L2/L3 \u5047\u8bbe\u8def\u5f84"],
        ["Stage 3", "GT \u6784\u5efa", "\u62bd\u53d6\u7ed3\u679c", "Benchmark Ground Truth"],
        ["Stage 4", "P1-P5 \u8bc4\u4f30", "\u8bba\u6587+\u63d0\u793a\u6a21\u677f", "\u591a\u7ef4\u5ea6\u8bc4\u5206"],
        ["Stage 5", "\u7efc\u5408\u5206\u6790", "P1-P5+KG\u6307\u6807", "\u8bc4\u4f30\u62a5\u544a"],
    ],
    col_widths=[2, 3, 4, 4],
)

# ══════════════════════════════════════════════════════════════════════
# Section 2: Example Paper
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("2. \u793a\u4f8b\u8bba\u6587 (Example Paper)", level=1)

title = paper0.get("title", "")
abstract = paper0.get("abstract", "")
primary = paper0.get("primary", "")
secondary = paper0.get("secondary", "")
url = paper0.get("pdf_url", "")

add_card_box("INPUT: \u8bba\u6587\u4fe1\u606f", "")

add_bold_text("Title: ", title)
add_bold_text("URL: ", url)
add_bold_text("Abstract: ", abstract[:500] + ("..." if len(abstract) > 500 else ""))

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# Section 3: Discipline Classification
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("3. Stage 1: \u5b66\u79d1\u5206\u7c7b (Discipline Classification)", level=1)

add_card_box("INPUT", f"Title: {title}\nAbstract: {abstract[:200]}...")
doc.add_paragraph()
add_card_box("OUTPUT", "", RGBColor(0x0D, 0x92, 0x76))

add_table(
    ["\u5b57\u6bb5", "\u503c"],
    [
        ["\u4e3b\u5b66\u79d1 (Primary)", primary],
        ["\u526f\u5b66\u79d1 (Secondary)", secondary],
    ],
    col_widths=[4, 9],
)

doc.add_paragraph(
    f"\u2192 \u8be5\u8bba\u6587\u88ab\u8bc6\u522b\u4e3a\u8de8\u5b66\u79d1\u8bba\u6587\uff0c\u4e3b\u5b66\u79d1\u4e3a\u300c{primary}\u300d\uff0c"
    f"\u6d89\u53ca {len(paper0.get('secondary_list', []))} \u4e2a\u5b66\u79d1\u9886\u57df\u3002"
)

# ══════════════════════════════════════════════════════════════════════
# Section 4: Concept Extraction
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("4. Stage 2a: \u6982\u5ff5\u62bd\u53d6 (Concept Extraction)", level=1)

concepts = paper0.get("concepts", {})
total_concepts = 0
concept_rows = []
for disc, clist in concepts.items():
    for c in clist:
        concept_rows.append([
            c.get("term", ""),
            c.get("normalized", ""),
            disc,
            str(c.get("confidence", "")),
        ])
        total_concepts += 1

add_card_box("INPUT", f"\u8bba\u6587\u5168\u6587 + \u5b66\u79d1\u5206\u7c7b\u7ed3\u679c (\u4e3b: {primary}, \u526f: {secondary})")
doc.add_paragraph()
add_card_box("OUTPUT", f"\u5171\u62bd\u53d6 {total_concepts} \u4e2a\u6982\u5ff5", RGBColor(0x0D, 0x92, 0x76))

# Show first 10 concepts
add_table(
    ["Term", "Normalized", "Discipline", "Confidence"],
    concept_rows[:12],
    col_widths=[5, 3, 3, 2],
)

if total_concepts > 12:
    doc.add_paragraph(f"  ... \u5171 {total_concepts} \u4e2a\u6982\u5ff5\uff0c\u4ec5\u5c55\u793a\u524d 12 \u4e2a")

# ══════════════════════════════════════════════════════════════════════
# Section 5: Relation Extraction
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("5. Stage 2b: \u5173\u7cfb\u62bd\u53d6 (Relation Extraction)", level=1)

relations = paper0.get("relations", [])

add_card_box("INPUT", f"{total_concepts} \u4e2a\u8de8\u5b66\u79d1\u6982\u5ff5")
doc.add_paragraph()
add_card_box("OUTPUT", f"\u5171\u62bd\u53d6 {len(relations)} \u6761\u5173\u7cfb", RGBColor(0x0D, 0x92, 0x76))

rel_rows = []
for r in relations:
    rel_rows.append([
        r.get("head", ""),
        r.get("relation_type", ""),
        r.get("tail", ""),
        str(r.get("confidence", "")),
    ])

add_table(
    ["Head", "Relation Type", "Tail", "Confidence"],
    rel_rows,
    col_widths=[4, 3, 4, 2],
)

# Show relation description
doc.add_heading("\u5173\u7cfb\u8be6\u60c5", level=2)
for r in relations[:3]:
    add_bold_text(
        f"{r.get('head', '')} \u2192 {r.get('tail', '')}:  ",
        r.get("relation", "")
    )

# ══════════════════════════════════════════════════════════════════════
# Section 6: Query Generation
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("6. Stage 2c: \u67e5\u8be2\u751f\u6210 (Query Generation)", level=1)

queries = paper0.get("queries", {})
l1_q = queries.get("L1", "")
l2_q = queries.get("L2", [])
l3_q = queries.get("L3", [])

add_card_box("INPUT", "\u6982\u5ff5 + \u5173\u7cfb + \u5b66\u79d1\u5206\u7c7b")
doc.add_paragraph()
add_card_box("OUTPUT", "\u4e09\u5c42\u6b21\u67e5\u8be2\u95ee\u9898", RGBColor(0x0D, 0x92, 0x76))

add_bold_text("L1 (\u5b8f\u89c2\u67e5\u8be2): ", l1_q if isinstance(l1_q, str) else str(l1_q))
if isinstance(l2_q, list):
    for i, q in enumerate(l2_q):
        add_bold_text(f"L2-{i+1} (\u4e2d\u89c2\u67e5\u8be2): ", q)
else:
    add_bold_text("L2 (\u4e2d\u89c2\u67e5\u8be2): ", str(l2_q))
if isinstance(l3_q, list):
    for i, q in enumerate(l3_q):
        add_bold_text(f"L3-{i+1} (\u5fae\u89c2\u67e5\u8be2): ", q)
else:
    add_bold_text("L3 (\u5fae\u89c2\u67e5\u8be2): ", str(l3_q))

# ══════════════════════════════════════════════════════════════════════
# Section 7: Hypothesis Path Generation
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("7. Stage 2d: \u5047\u8bbe\u8def\u5f84\u751f\u6210 (Hypothesis Path Generation)", level=1)

hyp_paths = paper0.get("hypothesis_paths", {})

add_card_box("INPUT", "\u67e5\u8be2\u95ee\u9898 + KG \u7ed3\u6784")
doc.add_paragraph()

total_paths = sum(len(v) if isinstance(v, list) else 0 for v in hyp_paths.values())
add_card_box("OUTPUT", f"\u5171\u751f\u6210 {total_paths} \u6761\u5047\u8bbe\u8def\u5f84 (L1/L2/L3)", RGBColor(0x0D, 0x92, 0x76))

for level_key in ["L1", "L2", "L3"]:
    paths = hyp_paths.get(level_key, [])
    if not paths:
        continue
    doc.add_heading(f"{level_key} Hypothesis Paths ({len(paths)} \u6761)", level=2)
    for pi, path in enumerate(paths[:2]):
        doc.add_paragraph(f"\u8def\u5f84 {pi+1}:", style="List Bullet")
        steps = path.get("steps", path.get("path", []))
        if isinstance(steps, list):
            for si, step in enumerate(steps):
                if isinstance(step, dict):
                    head = step.get("head", "")
                    tail = step.get("tail", "")
                    rel = step.get("relation", "")
                    claim = step.get("claim", "")
                    add_bold_text(f"  Step {si+1}: {head} \u2192 {tail}", f"  [{rel}]")
                    if claim:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(2)
                        run = p.add_run(f"\u2514 {claim[:200]}")
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ══════════════════════════════════════════════════════════════════════
# Section 8: KG Construction & Metrics
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("8. \u77e5\u8bc6\u56fe\u8c31\u6784\u5efa\u4e0e\u6307\u6807 (KG Construction & Metrics)", level=1)

add_card_box("INPUT", "\u6982\u5ff5 + \u5173\u7cfb + \u5047\u8bbe\u8def\u5f84")
doc.add_paragraph()

# Compute stats
n_papers = len(extraction)
all_concepts = [sum(len(v) for v in p.get("concepts", {}).values()) for p in extraction]
all_relations = [len(p.get("relations", [])) for p in extraction]
avg_c = sum(all_concepts) / max(len(all_concepts), 1)
avg_r = sum(all_relations) / max(len(all_relations), 1)

add_card_box("OUTPUT", f"\u56fe\u8c31\u7edf\u8ba1\u6982\u89c8", RGBColor(0x0D, 0x92, 0x76))

add_table(
    ["\u6307\u6807", "\u503c"],
    [
        ["\u5904\u7406\u8bba\u6587\u6570", str(n_papers)],
        ["\u5e73\u5747\u6982\u5ff5\u6570/\u7bc7", f"{avg_c:.1f}"],
        ["\u5e73\u5747\u5173\u7cfb\u6570/\u7bc7", f"{avg_r:.1f}"],
        ["\u5173\u7cfb\u7c7b\u578b\u6570", "12"],
        ["\u603b\u5173\u7cfb\u6570", str(sum(all_relations))],
    ],
    col_widths=[5, 5],
)

if kg0:
    doc.add_heading("KG \u8bc4\u4f30\u6307\u6807 (\u793a\u4f8b\u8bba\u6587)", level=2)
    scores = kg0.get("scores", {})
    kg_rows = []
    for level in ["L1", "L2", "L3"]:
        bridging = scores.get(f"{level}_bridging", "N/A")
        coherence = scores.get(f"{level}_chain_coherence", "N/A")
        novelty = scores.get(f"{level}_info_novelty", "N/A")
        innovation = scores.get(f"{level}_innovation", "N/A")
        feasibility = scores.get(f"{level}_feasibility", "N/A")
        entity_cov = scores.get(f"{level}_entity_coverage", "N/A")
        kg_rows.append([
            level,
            f"{bridging:.2f}" if isinstance(bridging, float) else str(bridging),
            f"{coherence:.3f}" if isinstance(coherence, float) else str(coherence),
            f"{novelty:.2f}" if isinstance(novelty, float) else str(novelty),
            f"{innovation:.1f}" if isinstance(innovation, float) else str(innovation),
            f"{feasibility:.1f}" if isinstance(feasibility, float) else str(feasibility),
            f"{entity_cov:.2f}" if isinstance(entity_cov, float) else str(entity_cov),
        ])
    add_table(
        ["Level", "Bridging", "Coherence", "Novelty", "Innovation", "Feasibility", "Coverage"],
        kg_rows,
        col_widths=[1.5, 2, 2, 2, 2, 2, 2],
    )

    # Depth analysis
    doc.add_heading("\u6df1\u5ea6\u5206\u6790 (Depth Analysis)", level=2)
    depth_rows = [
        ["L2 Concept Expansion", f"{scores.get('depth_l2_concept_expansion', 'N/A')}"],
        ["L3 Concept Expansion", f"{scores.get('depth_l3_concept_expansion', 'N/A')}"],
        ["L2 Anchoring", f"{scores.get('depth_l2_anchoring', 'N/A')}"],
        ["L3 Anchoring", f"{scores.get('depth_l3_anchoring', 'N/A')}"],
        ["Depth Quality", f"{scores.get('depth_depth_quality', 'N/A')}"],
    ]
    add_table(["\u6307\u6807", "\u503c"], depth_rows, col_widths=[5, 5])

# ══════════════════════════════════════════════════════════════════════
# Section 9: GT Construction
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("9. Benchmark GT \u6784\u5efa (Ground Truth Construction)", level=1)

gt_input = gt0.get("input", {})
gt_data = gt0.get("ground_truth", {})
gt_stats_data = gt0.get("gt_stats", {})

add_card_box("INPUT", f"\u62bd\u53d6\u7ed3\u679c: {total_concepts} \u4e2a\u6982\u5ff5, {len(relations)} \u6761\u5173\u7cfb, {total_paths} \u6761\u5047\u8bbe\u8def\u5f84")
doc.add_paragraph()
add_card_box("OUTPUT", "Evidence-Grounded Ground Truth", RGBColor(0x0D, 0x92, 0x76))

# GT stats overview
n_gt = len(benchmark)
all_gt_terms = [p.get("gt_stats", {}).get("n_terms", 0) for p in benchmark]
all_gt_rels = [p.get("gt_stats", {}).get("n_relations", 0) for p in benchmark]
avg_gt_terms = sum(all_gt_terms) / max(len(all_gt_terms), 1)
avg_gt_rels = sum(all_gt_rels) / max(len(all_gt_rels), 1)

add_table(
    ["\u6307\u6807", "\u503c"],
    [
        ["GT \u6761\u76ee\u6570", str(n_gt)],
        ["\u5e73\u5747\u672f\u8bed\u6570/\u6761", f"{avg_gt_terms:.1f}"],
        ["\u5e73\u5747\u5173\u7cfb\u6570/\u6761", f"{avg_gt_rels:.1f}"],
        ["\u603b\u672f\u8bed\u6570", str(sum(all_gt_terms))],
        ["\u603b\u5173\u7cfb\u6570", str(sum(all_gt_rels))],
    ],
    col_widths=[5, 5],
)

# Show example GT terms
doc.add_heading("\u793a\u4f8b GT \u672f\u8bed (\u524d 8 \u4e2a)", level=2)
gt_terms = gt_data.get("terms", [])
gt_term_rows = []
for t in gt_terms[:8]:
    gt_term_rows.append([
        t.get("term", ""),
        t.get("normalized", ""),
        t.get("discipline", ""),
        str(t.get("confidence", "")),
        str(t.get("grounded_to", "") or "-"),
    ])
add_table(
    ["Term", "Normalized", "Discipline", "Conf.", "Grounded To"],
    gt_term_rows,
    col_widths=[4, 3, 2, 1.5, 3],
)

# Show GT relations
doc.add_heading("\u793a\u4f8b GT \u5173\u7cfb", level=2)
gt_rels = gt_data.get("relations", [])
gt_rel_rows = []
for r in gt_rels[:5]:
    gt_rel_rows.append([
        r.get("head", ""),
        r.get("relation_type", ""),
        r.get("tail", ""),
        str(r.get("confidence", "")),
    ])
add_table(
    ["Head", "Type", "Tail", "Confidence"],
    gt_rel_rows,
    col_widths=[4, 3, 4, 2],
)

# ══════════════════════════════════════════════════════════════════════
# Section 10: P1-P5 Evaluation
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("10. P1-P5 \u8bc4\u4f30\u6846\u67b6 (Progressive Prompt Evaluation)", level=1)

doc.add_paragraph(
    "\u6e10\u8fdb\u5f0f\u63d0\u793a\u8bc4\u4f30\u6846\u67b6\u901a\u8fc7 5 \u4e2a\u7ea7\u522b\u7684\u4fe1\u606f\u63d0\u793a\u91cf\uff0c"
    "\u8861\u91cf LLM \u5728\u4e0d\u540c\u4fe1\u606f\u5145\u5206\u5ea6\u4e0b\u7684\u5047\u8bbe\u751f\u6210\u8d28\u91cf\u3002"
)

add_table(
    ["\u7ea7\u522b", "\u63cf\u8ff0"],
    [
        ["P1", "\u4ec5\u63d0\u4f9b\u8bba\u6587\u6807\u9898 (Title only)"],
        ["P2", "\u6807\u9898 + \u6458\u8981 (Title + Abstract)"],
        ["P3", "\u6807\u9898 + \u6458\u8981 + \u5b66\u79d1\u5206\u7c7b (+ Discipline)"],
        ["P4", "\u6807\u9898 + \u6458\u8981 + \u5b66\u79d1 + \u6982\u5ff5\u5217\u8868 (+ Concepts)"],
        ["P5", "\u5b8c\u6574\u77e5\u8bc6\u56fe\u8c31 + \u5047\u8bbe\u8def\u5f84 (Full KG + Paths)"],
    ],
    col_widths=[2, 11],
)

# P1-P5 comparison table
doc.add_heading("\u8bc4\u4f30\u7ed3\u679c\u5bf9\u6bd4", level=2)

p_levels = ["P1", "P2", "P3", "P4", "P5"]
eval_rows = []
for pl in p_levels:
    d = p1p5.get(pl, {})
    eval_rows.append([
        pl,
        f"{d.get('text_bertscore_f1', 'N/A'):.4f}" if isinstance(d.get('text_bertscore_f1'), float) else "N/A",
        f"{d.get('text_rouge1', 'N/A'):.4f}" if isinstance(d.get('text_rouge1'), float) else "N/A",
        f"{d.get('judge_novelty', 'N/A')}" if d.get('judge_novelty') is not None else "N/A",
        f"{d.get('judge_feasibility', 'N/A')}" if d.get('judge_feasibility') is not None else "N/A",
        f"{d.get('judge_relevance', 'N/A')}" if d.get('judge_relevance') is not None else "N/A",
        f"{d.get('judge_cross_disciplinary', 'N/A')}" if d.get('judge_cross_disciplinary') is not None else "N/A",
        f"{d.get('elapsed_seconds', 'N/A'):.1f}s" if isinstance(d.get('elapsed_seconds'), float) else "N/A",
    ])

add_table(
    ["Level", "BERTScore", "ROUGE-1", "Novelty", "Feasibility", "Relevance", "Cross-Disc", "Time"],
    eval_rows,
    col_widths=[1.5, 2, 2, 1.8, 1.8, 1.8, 1.8, 1.5],
)

# Key findings
doc.add_heading("\u5173\u952e\u53d1\u73b0", level=2)
add_bullet("P4 \u5728\u6587\u672c\u76f8\u4f3c\u5ea6 (BERTScore F1) \u4e0a\u8868\u73b0\u6700\u4f18: 0.7472")
add_bullet("P1 \u5728\u65b0\u9896\u6027 (Novelty=9.0) \u548c\u7279\u5f02\u6027 (Specificity=9.0) \u4e0a\u8868\u73b0\u6700\u4f18")
add_bullet("P5 \u5728\u53ef\u884c\u6027 (Feasibility=8.6) \u548c\u76f8\u5173\u6027 (Relevance=10.0) \u4e0a\u8868\u73b0\u6700\u4f18")
add_bullet("P5 \u751f\u6210\u901f\u5ea6\u6700\u5feb (\u53d6\u81ea\u77e5\u8bc6\u56fe\u8c31)\uff0c\u5e73\u5747\u751f\u6210 6.2 \u6761\u5047\u8bbe")
add_bullet("\u4fe1\u606f\u5145\u5206\u5ea6\u4e0e\u65b0\u9896\u6027\u5448\u8d1f\u76f8\u5173\uff0c\u4e0e\u76f8\u5173\u6027/\u53ef\u884c\u6027\u5448\u6b63\u76f8\u5173")

# ══════════════════════════════════════════════════════════════════════
# Section 11: Depth Analysis
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("11. \u6df1\u5ea6\u5206\u6790 (Depth Analysis: L1 vs L2 vs L3)", level=1)

doc.add_paragraph(
    "\u4e09\u5c42\u5047\u8bbe\u8def\u5f84\u4ee3\u8868\u4e0d\u540c\u62bd\u8c61\u7c92\u5ea6\u7684\u7814\u7a76\u95ee\u9898:\n"
    "\u2022 L1 (\u5b8f\u89c2): \u5b66\u79d1\u95f4\u5b8f\u89c2\u8054\u7cfb\uff0c\u5c11\u91cf\u6982\u5ff5\uff0c\u9ad8\u521b\u65b0\u6027\n"
    "\u2022 L2 (\u4e2d\u89c2): \u7279\u5b9a\u673a\u5236\u7ea7\u5173\u8054\uff0c\u6982\u5ff5\u6269\u5c55 25%\n"
    "\u2022 L3 (\u5fae\u89c2): \u7ec6\u7c92\u5ea6\u5b9e\u9a8c\u7ea7\u5173\u8054\uff0c\u6982\u5ff5\u6269\u5c55 62.5%"
)

if kg0:
    scores = kg0.get("scores", {})
    depth_table_rows = [
        ["Innovation",
         f"{scores.get('L1_innovation', 'N/A')}",
         f"{scores.get('L2_innovation', 'N/A')}",
         f"{scores.get('L3_innovation', 'N/A')}"],
        ["Feasibility",
         f"{scores.get('L1_feasibility', 'N/A')}",
         f"{scores.get('L2_feasibility', 'N/A')}",
         f"{scores.get('L3_feasibility', 'N/A')}"],
        ["Scientificity",
         f"{scores.get('L1_scientificity', 'N/A')}",
         f"{scores.get('L2_scientificity', 'N/A')}",
         f"{scores.get('L3_scientificity', 'N/A')}"],
        ["Testability",
         f"{scores.get('L1_testability', 'N/A')}",
         f"{scores.get('L2_testability', 'N/A')}",
         f"{scores.get('L3_testability', 'N/A')}"],
        ["Chain Coherence",
         f"{scores.get('L1_chain_coherence', 'N/A'):.3f}" if isinstance(scores.get('L1_chain_coherence'), float) else "N/A",
         f"{scores.get('L2_chain_coherence', 'N/A'):.3f}" if isinstance(scores.get('L2_chain_coherence'), float) else "N/A",
         f"{scores.get('L3_chain_coherence', 'N/A'):.3f}" if isinstance(scores.get('L3_chain_coherence'), float) else "N/A"],
        ["Entity Coverage",
         f"{scores.get('L1_entity_coverage', 'N/A')}",
         f"{scores.get('L2_entity_coverage', 'N/A')}",
         f"{scores.get('L3_entity_coverage', 'N/A')}"],
    ]
    add_table(
        ["\u6307\u6807", "L1", "L2", "L3"],
        depth_table_rows,
        col_widths=[4, 3, 3, 3],
    )

# ══════════════════════════════════════════════════════════════════════
# Section 12: Dataset Statistics Summary
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("12. \u6570\u636e\u96c6\u7edf\u8ba1\u6982\u89c8 (Dataset Statistics)", level=1)

add_table(
    ["\u7ef4\u5ea6", "\u7edf\u8ba1\u503c"],
    [
        ["\u8bba\u6587\u6570\u91cf", str(n_papers)],
        ["GT \u6761\u76ee\u6570", str(n_gt)],
        ["\u5e73\u5747\u6982\u5ff5/\u7bc7", f"{avg_c:.1f}"],
        ["\u5e73\u5747\u5173\u7cfb/\u7bc7", f"{avg_r:.1f}"],
        ["\u5e73\u5747 GT \u672f\u8bed/\u6761", f"{avg_gt_terms:.1f}"],
        ["\u5e73\u5747 GT \u5173\u7cfb/\u6761", f"{avg_gt_rels:.1f}"],
        ["\u603b\u672f\u8bed\u6570", str(sum(all_gt_terms))],
        ["\u603b\u5173\u7cfb\u6570", str(sum(all_gt_rels))],
        ["\u5b66\u79d1\u8986\u76d6\u6570", str(gt_stats_data.get("n_disciplines", "N/A"))],
        ["\u8de8\u5b66\u79d1\u8def\u5f84\u6570", str(gt_stats_data.get("n_cross_disciplinary_paths", "N/A"))],
    ],
    col_widths=[5, 5],
)

# ── Save ──
output_path = os.path.join(BASE, "CrossDisc_Bench_pipeline_walkthrough.docx")
doc.save(output_path)
print(f"Word document saved to: {output_path}")
