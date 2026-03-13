"""生成 Benchmark 项目运行流程详解 Word 文档."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    hs.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    hs.paragraph_format.space_after = Pt(6)

# ── 辅助函数 ──────────────────────────────────────────────────────────

def add_code_block(text: str):
    """添加代码块（灰底等宽字体）."""
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        # 灰色底纹
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        run.element.rPr.append(shd)


def add_json_block(obj, indent=2):
    """添加 JSON 代码块."""
    if isinstance(obj, str):
        add_code_block(obj)
    else:
        add_code_block(json.dumps(obj, ensure_ascii=False, indent=indent))


def add_table(headers, rows, col_widths=None):
    """添加表格."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表后空行


def add_bullet(text, level=0):
    """添加项目符号."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)


def add_bold_text(bold_part, normal_part):
    """添加加粗前缀 + 正常文本."""
    p = doc.add_paragraph()
    run_b = p.add_run(bold_part)
    run_b.bold = True
    p.add_run(normal_part)

# ══════════════════════════════════════════════════════════════════════
#  正文开始
# ══════════════════════════════════════════════════════════════════════

# ── 封面 ──
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("跨学科知识抽取与基准评测系统")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Cross-Disciplinary Knowledge Extraction & Benchmarking")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub2.add_run("项目运行流程与输入输出详解  ·  v0.2.0")
run3.font.size = Pt(16)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 第一章：项目概述
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("一、项目概述", level=1)

doc.add_paragraph(
    "本项目是一个基于多阶段大语言模型（LLM）Pipeline 的跨学科知识抽取与基准评测系统。"
    "系统从学术论文中自动抽取跨学科概念、关系，生成科学假设，并通过多维度指标进行评估。"
)

doc.add_heading("1.1 核心能力", level=2)
add_bullet("学科分类：基于 MSC 层级分类体系，自动识别论文的主学科与辅学科")
add_bullet("概念抽取：从论文标题、摘要、引言中抽取专业术语，按学科分组")
add_bullet("关系抽取：识别跨学科概念间的语义关系（11 种标准关系类型）")
add_bullet("查询生成：生成三级层次化查询（宏观 → 学科维度 → 细粒度）")
add_bullet("假设生成：生成三级知识路径（每条路径 3 步链式推理）")
add_bullet("知识图谱构建：基于 NetworkX 构建概念图谱并计算 15+ 图谱指标")
add_bullet("基准评测：LLM-as-Judge 主观评分 + 客观图谱指标的多维度评估")

doc.add_heading("1.2 技术栈", level=2)
add_table(
    ["组件", "技术选型", "用途"],
    [
        ["LLM 调用", "OpenAI 兼容 API (qwen3-235b-a22b)", "概念/关系/假设抽取 + 评估"],
        ["数据模型", "Pydantic v2", "严格 JSON Schema 验证"],
        ["图谱构建", "NetworkX", "概念图谱 + 拓扑指标"],
        ["重试机制", "tenacity", "指数退避重试（最多 5 次）"],
        ["分类器", "LangChain + OpenAI", "层级学科分类"],
        ["PDF 处理", "pdfminer / pdfplumber", "论文引言提取"],
        ["语义相似度", "sentence-transformers (可选)", "嵌入桥接 + 链式连贯性"],
        ["配置管理", "PyYAML + frozen dataclass", "线程安全配置"],
    ],
    col_widths=[3.5, 5.5, 7],
)

# ══════════════════════════════════════════════════════════════════════
# 第二章：端到端运行总览
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("二、端到端运行总览", level=1)

doc.add_paragraph(
    "系统完整运行分为 6 个阶段，从原始论文输入到最终多维度评分输出。"
    "以下为端到端流程图："
)

add_code_block("""┌─────────────────────────────────────────────────────────────────────┐
│  crossdisc-pipeline full -i papers.jsonl -o results.jsonl          │
│                                                                     │
│  Phase 0: 数据准备 (输入论文 JSONL)                                  │
│      ↓                                                              │
│  Phase 1: 学科分类 (classifier/) → 主学科 + 辅学科                   │
│      ↓                                                              │
│  Phase 2: 多学科过滤 → 仅保留跨学科论文                              │
│      ↓                                                              │
│  Phase 3: 三阶段知识抽取 (extractor_multi_stage.py)                  │
│      ├─ Stage 1a: 概念抽取 (LLM Round 1 + Round 2 补充)             │
│      ├─ Stage 1b: 关系抽取                                          │
│      ├─ Stage 2:  查询生成 (三级)                                    │
│      └─ Stage 3:  假设生成 (L1/L2/L3)                               │
│      ↓                                                              │
│  Phase 4: 图谱构建 + 指标计算 (graph_builder.py)                     │
│      ↓                                                              │
│  Phase 5: Ground Truth 构建 (gt_builder.py)                         │
│      ↓                                                              │
│  Phase 6: 基准评测 (evaluate_benchmark.py) → 多维度评分              │
└─────────────────────────────────────────────────────────────────────┘""")

doc.add_heading("2.1 LLM 调用次数统计（单篇论文）", level=2)
add_table(
    ["阶段", "LLM 调用次数", "说明"],
    [
        ["Phase 1 分类", "2-4 次", "每层分类体系一次（通常 2-3 层）"],
        ["Stage 1a 概念", "2 次", "Round 1 + Round 2 补充"],
        ["Stage 1b 关系", "1 次", "基于概念列表抽取关系"],
        ["Stage 2 查询", "1 次", "生成三级查询"],
        ["Stage 3 假设", "3 次", "L1 + L2 + L3 各一次"],
        ["总计", "~9-11 次", "单篇论文的完整处理"],
    ],
    col_widths=[4, 3, 9],
)

# ══════════════════════════════════════════════════════════════════════
# 第三章：Phase 0 — 数据准备
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("三、Phase 0：数据准备", level=1)

doc.add_heading("3.1 CLI 命令", level=2)
doc.add_paragraph("推荐使用统一 Pipeline 命令：")
add_code_block("crossdisc-pipeline full -i papers.jsonl -o results.jsonl --config configs/default.yaml")

doc.add_paragraph("入口文件：pipeline.py → main() → argparse 解析子命令 full")

doc.add_heading("3.2 输入格式", level=2)
doc.add_paragraph("支持 .json / .jsonl / .csv 三种格式，字段名兼容中英文变体（如 标题/title、摘要/abstract）。")
doc.add_paragraph("JSONL 格式示例（每行一篇论文）：")
add_json_block({
    "title": "Refining centromedian nucleus stimulation for generalized epilepsy with targeting and mechanistic insights from intraoperative electrophysiology",
    "abstract": "Epilepsy affects 65 million people worldwide, with 30% suffering from drug-resistant epilepsy. While surgical resection is the primary treatment...",
    "pdf_url": "https://www.nature.com/articles/s41467-025-60183-9.pdf"
})

doc.add_heading("3.3 必需字段与可选字段", level=2)
add_table(
    ["字段", "是否必需", "说明"],
    [
        ["title / 标题", "必需", "论文标题"],
        ["abstract / 摘要", "必需", "论文摘要"],
        ["pdf_url", "可选", "PDF 链接，用于提取引言（Introduction）"],
        ["primary / 主学科", "可选*", "主学科（若已有则跳过分类阶段）"],
        ["secondary_list", "可选*", "辅学科列表（若已有则跳过分类阶段）"],
    ],
    col_widths=[4, 2.5, 9.5],
)
doc.add_paragraph("* 若输入已包含 primary 和 secondary_list 字段，可直接使用 extract 子命令跳过分类阶段。")

# ══════════════════════════════════════════════════════════════════════
# 第四章：Phase 1 — 学科分类
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("四、Phase 1：学科分类", level=1)

doc.add_paragraph(
    "入口：pipeline.py → classify_and_filter() → AsyncHierarchicalClassifier"
)

doc.add_heading("4.1 分类体系", level=2)
doc.add_paragraph(
    "系统使用 MSC（数学学科分类）层级分类体系，存储于 data/msc_converted.json。"
    "该分类体系为多层树状结构，例如："
)
add_code_block("""数学
├── 数论
│   ├── 初等数论
│   └── 解析数论
├── 数学分析
│   ├── 实分析
│   └── 复分析
临床医学
├── 内科学
├── 外科学
├── 神经病学
...""")

doc.add_heading("4.2 分类流程", level=2)
doc.add_paragraph("对每篇论文，分类器逐层遍历分类体系，每层调用一次 LLM：")
add_bullet("Level 0：从顶层学科中选择（数学、物理学、化学、临床医学、基础医学、工程与技术科学基础学科...）")
add_bullet("Level 1：在选中的 L0 学科下选择子类")
add_bullet("Level 2：继续细分（如有更深层级）")

doc.add_paragraph("每层 LLM 交互流程：")
add_bullet("构建 Prompt：包含论文标题 + 摘要 + 当前层候选选项 + 子选项展开（提供上下文）")
add_bullet("LLM 输出：方括号列表格式，如 [临床医学, 基础医学, 工程与技术科学基础学科]")
add_bullet("校验：验证输出选项是否在合法候选列表中，不合法则重试（最多 max_retries 次）")
add_bullet("多学科判定：当分类路径中涉及多于 1 个不同的一级学科（L1）时标记为跨学科论文")

doc.add_heading("4.3 输入 → 输出示例", level=2)
doc.add_paragraph("输入：")
add_json_block({
    "title": "Refining centromedian nucleus stimulation for generalized epilepsy...",
    "abstract": "Epilepsy affects 65 million people worldwide..."
})

doc.add_paragraph("LLM 交互（Level 0）：")
add_code_block("""Prompt: "请从以下学科中选择该论文涉及的学科: [数学, 物理学, 化学, 临床医学, 基础医学, 工程与技术科学基础学科, ...]"
LLM 输出: "[临床医学, 基础医学, 工程与技术科学基础学科]" """)

doc.add_paragraph("输出：")
add_json_block({
    "title": "Refining centromedian nucleus stimulation...",
    "abstract": "Epilepsy affects 65 million...",
    "pdf_url": "https://www.nature.com/articles/s41467-025-60183-9.pdf",
    "primary": "临床医学",
    "secondary": "基础医学, 工程与技术科学基础学科",
    "secondary_list": ["基础医学", "工程与技术科学基础学科"],
    "is_multidisciplinary": True
})

# ══════════════════════════════════════════════════════════════════════
# 第五章：Phase 2 — 多学科过滤
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("五、Phase 2：多学科过滤", level=1)
doc.add_paragraph(
    "pipeline.py 中 full 子命令在分类完成后执行过滤逻辑："
)
add_bullet("检查分类结果中 secondary_list 是否非空")
add_bullet("仅保留跨学科论文（secondary_list 非空）进入后续抽取阶段")
add_bullet("单学科论文被过滤掉，不进入知识抽取流程")
doc.add_paragraph(
    "过滤的目的是确保后续的跨学科关系抽取和假设生成有意义——"
    "只有涉及多个学科的论文才具备跨学科知识桥接的分析价值。"
)

# ══════════════════════════════════════════════════════════════════════
# 第六章：Phase 3 — 三阶段知识抽取
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("六、Phase 3：三阶段知识抽取", level=1)

doc.add_paragraph(
    "入口：extractor_multi_stage.py → run_benchmark() → run_pipeline_for_item()\n"
    "这是系统的核心阶段，包含 Stage 1（结构抽取）、Stage 2（查询生成）、Stage 3（假设生成）三个子阶段。"
)

# ── Stage 1a ──
doc.add_heading("6.1 Stage 1a：概念抽取（Round 1）", level=2)
doc.add_paragraph("调用链：build_concepts_messages() → LLM → parse_concepts_output()")
doc.add_paragraph("LLM System Prompt 核心指令：")
add_bullet("从标题/摘要/引言中抽取专业术语和专有名词")
add_bullet("按主学科和辅学科分组")
add_bullet("每个术语需提供：term（原始术语）、normalized（标准化形式）、evidence（原文证据 ≤40 中文字）、source、confidence")

doc.add_paragraph("输入（User Prompt 变量）：")
add_code_block("""title = "Refining centromedian nucleus stimulation..."
abstract = "Epilepsy affects 65 million..."
introduction = "[从 PDF 提取的引言文本]"
primary = "临床医学"
secondary_list = "基础医学, 工程与技术科学基础学科" """)

doc.add_paragraph("输出示例：")
add_json_block({
    "meta": {
        "title": "Refining centromedian nucleus stimulation...",
        "primary": "临床医学",
        "secondary_list": ["基础医学", "工程与技术科学基础学科"]
    },
    "概念": {
        "主学科": [
            {"term": "癫痫", "normalized": "Epilepsy", "std_label": "Epilepsy",
             "evidence": "Epilepsy affects 65 million people worldwide",
             "source": "abstract", "confidence": 1.0},
            {"term": "脑深部刺激", "normalized": "Deep brain stimulation",
             "evidence": "Centromedian nucleus neurostimulation offers a promising alternative",
             "source": "abstract", "confidence": 1.0}
        ],
        "辅学科": {
            "基础医学": [
                {"term": "神经学", "normalized": "Neurology", "confidence": 0.9},
                {"term": "神经生理学", "normalized": "Neurophysiology", "confidence": 0.9},
                {"term": "电生理学", "normalized": "Electrophysiology", "confidence": 1.0},
                {"term": "神经回路", "normalized": "Neural circuits", "confidence": 0.8}
            ],
            "工程与技术科学基础学科": [
                {"term": "影像引导技术", "normalized": "Imaging-guided technology", "confidence": 0.8},
                {"term": "信号处理", "normalized": "Signal processing", "confidence": 0.8},
                {"term": "神经工程", "normalized": "Neural engineering", "confidence": 0.8}
            ]
        }
    }
})

# ── Stage 1a 补充 ──
doc.add_heading("6.2 Stage 1a 补充：概念补充抽取（Round 2）", level=2)
add_bullet("将 Round 1 已抽取的概念列表传入 LLM，要求补充遗漏的术语")
add_bullet("通过 _merge_concepts() 去重合并两轮结果")
add_bullet("通过 _ground_and_filter_concepts() 过滤学科名/通用词，并与 MSC 术语库对齐")
doc.add_paragraph("两轮抽取的设计目的是提高概念覆盖率，避免单次 LLM 调用遗漏重要术语。")

# ── Stage 1b ──
doc.add_heading("6.3 Stage 1b：关系抽取", level=2)
doc.add_paragraph("调用链：build_relations_messages() → LLM → parse_relations_output()")
doc.add_paragraph("输入：Stage 1a 的完整概念 JSON + 原文（标题 + 摘要 + 引言）")
doc.add_paragraph("输出示例：")
add_json_block({
    "跨学科关系": [
        {
            "head": "Deep brain stimulation",
            "relation": "offers a promising alternative treatment for",
            "relation_type": "method_applied_to",
            "tail": "Epilepsy",
            "direction": "->",
            "quant": None,
            "assumptions": ["Drug-resistant epilepsy is present"],
            "evidence": "Centromedian nucleus neurostimulation offers a promising alternative...",
            "source": "abstract",
            "confidence": 0.9
        },
        {
            "head": "Electrophysiology",
            "relation": "guides the optimization of",
            "relation_type": "improves_metric",
            "tail": "Deep brain stimulation",
            "evidence": "We present a multimodal approach integrating intraoperative thalamic...",
            "confidence": 0.9
        },
        {
            "head": "Signal processing",
            "relation": "provides real-time feedback for",
            "relation_type": "depends_on",
            "tail": "Electrophysiology",
            "evidence": "Simultaneous intraoperative EEG to deliver real-time feedback...",
            "confidence": 0.8
        }
    ]
})

doc.add_paragraph("系统支持 11 种标准关系类型：")
add_table(
    ["关系类型", "含义", "示例"],
    [
        ["method_applied_to", "方法应用于", "Deep brain stimulation → Epilepsy"],
        ["maps_to", "映射到", "Neurostimulation → Internal medullary lamina"],
        ["constrains", "约束", "Anatomy → Electrode placement"],
        ["improves_metric", "改善指标", "Electrophysiology → DBS optimization"],
        ["corresponds_to", "对应于", "Delta power → Seizure reduction"],
        ["inferred_from", "推断自", "Mechanism → Experimental data"],
        ["assumes", "假设", "Model → Drug resistance present"],
        ["extends", "扩展", "New method → Existing approach"],
        ["generalizes", "泛化", "Specific finding → General principle"],
        ["driven_by", "驱动于", "Outcome → Underlying mechanism"],
        ["depends_on", "依赖于", "Signal processing → Electrophysiology"],
    ],
    col_widths=[3.5, 2.5, 10],
)

doc.add_paragraph("Stage 1 最终输出 → StructExtraction 对象（meta + 概念 + 跨学科关系）")

# ── Stage 2 ──
doc.add_page_break()
doc.add_heading("6.4 Stage 2：查询生成", level=2)
doc.add_paragraph("调用链：build_query_messages() → LLM → parse_query_output()")
doc.add_paragraph("输入：Stage 1 的完整 StructExtraction JSON")
doc.add_paragraph("输出示例：")
add_json_block({
    "按辅助学科分类": {
        "基础医学": {
            "概念": ["神经学", "神经生理学", "神经解剖学", "神经调控", "电生理学", "神经回路"],
            "关系": [0, 1, 2, 3, 4, 5],
            "rationale": "基础医学在论文中主要提供神经机制、神经结构和电生理基础，为神经调控提供理论支持。"
        },
        "工程与技术科学基础学科": {
            "概念": ["影像引导技术", "电子工程", "信号处理", "神经工程"],
            "关系": [0, 1, 2, 3],
            "rationale": "工程与技术学科支撑高精度定位、信号采集与分析，推动神经调控技术的实际应用。"
        }
    },
    "查询": {
        "一级": "如何通过多学科技术优化脑深部刺激治疗癫痫？",
        "二级": [
            "结合基础医学中的神经生理学与神经解剖学",
            "利用工程科学的影像引导与信号处理技术"
        ],
        "三级": [
            "科研团队如何协同运用神经解剖学、神经生理学和信号处理技术实现精准靶向",
            "通过多模态成像引导技术与实时信号分析，验证脑深部刺激的机制与效果"
        ]
    }
})

doc.add_paragraph("查询层级含义：")
add_table(
    ["层级", "含义", "示例"],
    [
        ["一级 (L1)", "宏观/论文级别的核心问题", "如何通过多学科技术优化脑深部刺激治疗癫痫？"],
        ["二级 (L2)", "按辅助学科维度拆分的子问题", "结合基础医学中的神经生理学与神经解剖学"],
        ["三级 (L3)", "具体操作/方法层面的细粒度问题", "如何协同运用神经解剖学、信号处理技术实现精准靶向"],
    ],
    col_widths=[2.5, 5, 8.5],
)

# ── Stage 3 ──
doc.add_page_break()
doc.add_heading("6.5 Stage 3：假设生成（L1 / L2 / L3）", level=2)
doc.add_paragraph(
    "分三次独立调用 LLM，分别生成 L1（宏观）、L2（学科维度）、L3（细粒度）假设。"
    "每次调用输入为 StructExtraction 摘要 + Query3Levels。"
)

doc.add_heading("L1 输出示例（宏观假设 — 每条路径 3 步链式推理）", level=3)
add_json_block("""{
  "假设": {
    "一级": [
      [
        {"step": 1, "head": "脑深部刺激", "relation": "应用于", "tail": "癫痫治疗",
         "claim": "脑深部刺激是治疗药物难治性癫痫的一种有前景的方法。"},
        {"step": 2, "head": "癫痫治疗", "relation": "面临", "tail": "作用机制不明确",
         "claim": "当前脑深部刺激治疗癫痫的神经调控机制尚不完全清楚。"},
        {"step": 3, "head": "作用机制不明确", "relation": "需要", "tail": "基础医学补位",
         "claim": "需要基础医学通过神经生理学和神经解剖学揭示调控机制。"}
      ],
      [
        {"step": 1, "head": "脑深部刺激", "relation": "依赖于", "tail": "精准靶向",
         "claim": "脑深部刺激的治疗效果高度依赖于靶点的精准定位。"},
        {"step": 2, "head": "精准靶向", "relation": "受限于", "tail": "传统定位方法",
         "claim": "传统的解剖定位方法难以实现个体化、高精度的电极植入。"},
        {"step": 3, "head": "传统定位方法", "relation": "需要", "tail": "工程与技术科学基础学科补位",
         "claim": "需要工程与技术科学通过影像引导和信号处理技术提升精准度。"}
      ]
    ],
    "一级总结": [
      "脑深部刺激治疗癫痫需要基础医学揭示调控机制以优化刺激靶点",
      "精准靶向需要工程与技术科学提升电极植入精准度与实时反馈能力"
    ]
  }
}""")

doc.add_heading("L2 输出示例（按辅学科维度的假设）", level=3)
add_json_block("""{
  "假设": {
    "二级": [
      [
        {"step": 1, "head": "神经生理学", "relation": "提供癫痫发作的", "tail": "神经电生理信号",
         "claim": "神经生理学提供癫痫发作的神经电生理信号"},
        {"step": 2, "head": "神经电生理信号", "relation": "反映为刺激诱导的", "tail": "皮质δ功率变化",
         "claim": "神经电生理信号反映为刺激诱导的皮质δ功率变化"},
        {"step": 3, "head": "皮质δ功率变化", "relation": "作为优化", "tail": "脑深部刺激疗效",
         "claim": "皮质δ功率变化作为优化脑深部刺激疗效的关键指标"}
      ]
    ],
    "二级总结": ["神经生理学通过电生理信号指导脑深部刺激的靶点优化"]
  }
}""")

doc.add_heading("关键约束", level=3)
add_bullet("每条路径严格 3 步：step 1 → step 2 → step 3")
add_bullet("链式一致性：step[i].tail 必须与 step[i+1].head 语义匹配（相似度 ≥ 0.75）")
add_bullet("head/tail 必须来自已抽取的概念列表（不允许凭空创造）")
add_bullet("最后一步的 claim 必须非空，且为完整可验证的假设陈述")
doc.add_paragraph(
    "Stage 3 之后还会执行 _align_hypothesis_entities() 将假设中的实体与抽取概念对齐，"
    "确保知识图谱的节点一致性。"
)

# ══════════════════════════════════════════════════════════════════════
# 第七章：Phase 4 — 图谱构建 + 指标计算
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("七、Phase 4：图谱构建与指标计算", level=1)
doc.add_paragraph("入口：graph_builder.py → build_graph_and_metrics()")

doc.add_heading("7.1 构建流程", level=2)
add_bullet("添加节点：从 Stage 1 概念中提取所有术语作为图节点，标记学科归属")
add_bullet("添加结构边：从 Stage 1 关系中提取 head→tail 作为图边（带关系类型）")
add_bullet("添加假设边：从 Stage 3 每个 HypothesisStep 的 head→tail 作为图边")
add_bullet("计算指标：15+ 维度的图谱质量指标")

doc.add_heading("7.2 输出示例（ConceptGraph + GraphMetrics）", level=2)
add_json_block({
    "graph": {
        "nodes": [
            {"id": "癫痫", "discipline": "临床医学"},
            {"id": "脑深部刺激", "discipline": "临床医学"},
            {"id": "电生理学", "discipline": "基础医学"},
            {"id": "信号处理", "discipline": "工程与技术科学基础学科"}
        ],
        "edges": [
            {"source": "脑深部刺激", "target": "癫痫", "relation": "method_applied_to"},
            {"source": "电生理学", "target": "脑深部刺激", "relation": "improves_metric"}
        ]
    },
    "metrics": {
        "path_consistency": 0.85,
        "coverage": 0.72,
        "bridging_score": 0.90,
        "rao_stirling_diversity": 0.65,
        "embedding_bridging": 0.78,
        "chain_coherence": 0.82,
        "kg_density": 0.45,
        "kg_modularity": 0.38,
        "kg_betweenness": 0.22,
        "kg_clustering": 0.55,
        "atypical_combination": 0.31
    }
})

doc.add_heading("7.3 指标说明", level=2)
add_table(
    ["指标", "计算方法", "含义"],
    [
        ["path_consistency", "假设边与结构边的重叠率", "假设路径与原文关系的一致程度"],
        ["coverage", "假设覆盖的概念占比", "假设路径对抽取概念的覆盖广度"],
        ["bridging_score", "跨学科概念的 Jaccard 距离", "假设路径的跨学科桥接程度"],
        ["rao_stirling_diversity", "Δ = Σ d_ij · p_i · p_j", "学科多样性/均衡性/差异性"],
        ["embedding_bridging", "首尾概念的语义距离", "知识路径的语义跨度"],
        ["chain_coherence", "逐跳语义连贯性 (SBERT)", "推理链条的逻辑连贯程度"],
        ["kg_density", "图密度 = 2E / V(V-1)", "知识图谱的连接密度"],
        ["kg_modularity", "Newman 模块度", "图谱的社区结构清晰度"],
        ["kg_betweenness", "平均介数中心性", "关键桥接节点的重要性"],
        ["kg_clustering", "平均聚类系数", "局部连接的紧密程度"],
        ["atypical_combination", "Uzzi et al. z-score", "概念组合的新颖性"],
    ],
    col_widths=[3.5, 4.5, 8],
)

# ══════════════════════════════════════════════════════════════════════
# 第八章：Phase 5 — Ground Truth 构建
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("八、Phase 5：Ground Truth 构建", level=1)
doc.add_paragraph("入口：benchmark/gt_builder.py → build_ground_truth()")

doc.add_heading("8.1 三阶段构建流程", level=2)

add_bold_text("Stage A — 术语抽取：", "")
add_bullet("优先级：已解析概念（生产 Pipeline 输出）> LLM 抽取 > 启发式抽取")
add_bullet("启发式方法：缩写词识别、科学短语匹配、首字母大写短语、中文术语")
add_bullet("所有术语与 MSC 术语库进行相似度对齐（阈值 0.70）")

add_bold_text("Stage B — 关系构建：", "")
add_bullet("对原文分句 → 找术语共现对 → LLM 或启发式分类关系类型")
add_bullet("每条关系携带原文证据句（evidence_sentence）")

add_bold_text("Stage C — 路径构建：", "")
add_bullet("用 NetworkX 建有向图 → nx.all_simple_paths() 寻找跨学科路径")
add_bullet("路径最长 4 跳（cutoff=4）")
add_bullet("优先选择跨越多个学科的路径")

doc.add_heading("8.2 输出示例", level=2)
add_json_block({
    "id": "Refining centromedian nucleus stimulation...",
    "input": {
        "title": "Refining centromedian nucleus stimulation...",
        "primary_discipline": "临床医学",
        "secondary_disciplines": ["基础医学", "工程与技术科学基础学科"],
        "abstract": "Epilepsy affects 65 million..."
    },
    "ground_truth": {
        "terms": [
            {"term": "Neurological Surgery", "normalized": "neurological surgery",
             "source": "heuristic", "confidence": 0.4},
            {"term": "Neural Engineering Labs", "normalized": "neural engineering labs",
             "source": "heuristic", "confidence": 0.4}
        ],
        "relations": [
            {"head": "neural engineering", "tail": "electrophysiology",
             "relation_type": "depends_on",
             "evidence_sentence": "multimodal approach integrating intraoperative...",
             "confidence": 0.5, "source_method": "heuristic"}
        ],
        "paths": [],
        "concept_graph": {"nodes": ["..."], "edges": ["..."]}
    }
})

# ══════════════════════════════════════════════════════════════════════
# 第九章：Phase 6 — 基准评测
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("九、Phase 6：基准评测", level=1)
doc.add_paragraph("入口：benchmark/evaluate_benchmark.py")

doc.add_heading("9.1 评测流程", level=2)
add_bullet("加载 benchmark 数据集 → 构建 GlobalKG（全局知识图谱索引）")
add_bullet("对每篇论文的 L1/L2/L3 假设路径分别评估")
add_bullet("结合 LLM 主观评分与客观图谱指标，输出多维度评分")

doc.add_heading("9.2 评测维度", level=2)
add_table(
    ["评测维度", "方法", "分值范围", "说明"],
    [
        ["Innovation（创新性）", "LLM-as-Judge", "0-10", "假设的新颖程度和原创性"],
        ["Feasibility（可行性）", "LLM-as-Judge", "0-10", "假设的实验可行性"],
        ["Scientificity（科学性）", "LLM-as-Judge", "0-10", "假设的科学严谨性"],
        ["Consistency（一致性）", "链式一致性检查", "0-1", "推理链条的逻辑一致性"],
        ["Bridging（桥接度）", "跨学科概念距离", "0-1", "跨学科知识桥接程度"],
        ["Rao-Stirling Diversity", "学科多样性公式", "0-1", "涉及学科的多样性"],
        ["Chain Coherence", "SBERT 语义连贯性", "0-1", "逐跳推理的语义连贯"],
        ["Atypical Combination", "Uzzi et al. z-score", "连续值", "概念组合的非典型性"],
        ["KG Topology", "图论指标", "连续值", "图密度/模块度/介数/聚类"],
        ["Concept Coverage", "GT 术语匹配", "0-1", "召回率/精确率/F1"],
        ["Path Alignment", "嵌入相似度", "0-1", "与 GT 路径的语义对齐度"],
    ],
    col_widths=[3.5, 3, 2, 7.5],
)

doc.add_heading("9.3 输出示例", level=2)
add_json_block({
    "id": "-33484234247962345",
    "scores": {
        "L1_innovation": 7.35,
        "L1_feasibility": 7.80,
        "L1_scientificity": 8.70,
        "L1_consistency": 1.0,
        "L1_bridging": 1.0,
        "L2_innovation": 7.40,
        "L2_feasibility": 7.25,
        "L2_scientificity": 6.45,
        "L2_consistency": 0.5,
        "L2_bridging": 1.0,
        "L3_innovation": 8.05,
        "L3_feasibility": 5.25,
        "L3_scientificity": 6.70,
        "L3_consistency": 0.0,
        "L3_bridging": 1.0
    }
})

# ══════════════════════════════════════════════════════════════════════
# 第十章：CLI 命令参考与配置
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("十、CLI 命令参考与配置", level=1)

doc.add_heading("10.1 端到端 Pipeline（推荐）", level=2)
add_code_block("""# 完整流水线：分类 → 过滤 → 抽取
crossdisc-pipeline full \\
  -i papers.jsonl \\
  -o results.jsonl \\
  --config configs/default.yaml

# 仅分类
crossdisc-pipeline classify \\
  -i papers.jsonl \\
  -o classified.jsonl \\
  --config configs/default.yaml

# 仅抽取（已有分类结果）
crossdisc-pipeline extract \\
  -i classified.jsonl \\
  -o results.jsonl""")

doc.add_heading("10.2 独立三阶段抽取", level=2)
add_code_block("""# 批量抽取（输入需已包含 primary/secondary_list 字段）
python run.py batch \\
  --input data.jsonl \\
  --output out.jsonl \\
  --num-workers 4 \\
  --resume

# 单篇抽取
python run.py single \\
  --title "论文标题" \\
  --abstract "论文摘要" \\
  --primary "临床医学" \\
  --secondary "基础医学,工程与技术科学基础学科" """)

doc.add_heading("10.3 关键配置参数", level=2)
doc.add_paragraph("配置文件：configs/experiment_v1.yaml")
add_table(
    ["参数", "默认值", "说明"],
    [
        ["language_mode", "chinese", "语言模式：chinese | original"],
        ["temperature_struct", "0.2", "Stage 1 温度（低 → 准确）"],
        ["temperature_query", "0.2", "Stage 2 温度"],
        ["temperature_hyp", "0.3", "Stage 3 温度（略高 → 多样性）"],
        ["seed", "42", "LLM 采样种子（可复现性）"],
        ["num_workers", "1", "并行 worker 数（1=串行）"],
        ["resume", "true", "断点续传（仅 JSONL 输出有效）"],
        ["max_tokens_struct", "8192", "Stage 1 最大 token"],
        ["max_tokens_query", "4096", "Stage 2 最大 token"],
        ["max_tokens_hyp", "4096", "Stage 3 每个子阶段最大 token"],
    ],
    col_widths=[4, 2.5, 9.5],
)

doc.add_heading("10.4 断点续跑机制", level=2)
add_bullet("输出格式必须为 .jsonl（每处理完一篇论文立即追加一行）")
add_bullet("每篇论文用 MD5(title) 作为唯一 ID")
add_bullet("重启时读取已有输出 → 收集 ok=True 的记录 ID → 跳过已完成的论文")
add_bullet("支持串行和并行模式")

# ══════════════════════════════════════════════════════════════════════
# 第十一章：最终输出数据结构
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("十一、最终输出数据结构（Extraction）", level=1)
doc.add_paragraph("一篇论文处理完成后的完整 JSON 结构：")
add_code_block("""{
  "meta": {title, primary, secondary_list},
  "概念": {
    "主学科": [ConceptEntry...],
    "辅学科": {"学科名": [ConceptEntry...]}
  },
  "跨学科关系": [RelationEntry...],
  "按辅助学科分类": {
    "学科名": {"概念": [...], "关系": [...], "rationale": "..."}
  },
  "查询": {"一级": str, "二级": [str], "三级": [str]},
  "假设": {
    "一级": [[Step,Step,Step]...], "一级总结": [str...],
    "二级": [[Step,Step,Step]...], "二级总结": [str...],
    "三级": [[Step,Step,Step]...], "三级总结": [str...]
  },
  "graph": {"nodes": [...], "edges": [...]},
  "metrics": {path_consistency, coverage, bridging_score, ...15+指标}
}""")

doc.add_heading("数据模型层级", level=2)
add_code_block("""Extraction（完整抽取结果）
├── meta: {title, primary, secondary_list}
├── 概念 (Concepts): 主学科 + 辅学科概念
│   └── ConceptEntry: {term, normalized, std_label, evidence, source, confidence}
├── 跨学科关系: RelationEntry[]
│   └── RelationEntry: {head, tail, relation, relation_type, evidence, confidence}
├── 查询 (Query3Levels): L1/L2/L3
├── 假设 (Hypothesis3Levels): L1/L2/L3 路径 + 总结
│   └── HypothesisStep: {step, head, relation, tail, claim}
├── graph (ConceptGraph): nodes + edges
└── metrics (GraphMetrics): 15+ 指标""")

# ── 保存 ──
output_path = "/ssd/wangyuyang/git/benchmark/outputs/项目运行流程详解.docx"
doc.save(output_path)
print(f"Word 文档已保存至: {output_path}")
