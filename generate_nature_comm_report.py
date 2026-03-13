"""
generate_nature_comm_report.py
Nature Communications 100篇 全流程汇报文档生成器。

读取 outputs/nature_comm_100/stage_outputs/ 下各阶段样例，
生成一份适合向领导汇报的 Word 文档，每阶段展示 3 个真实样例。
"""
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 路径配置 ──────────────────────────────────────────────────────────
OUTPUT_DIR = os.environ.get("REPORT_OUTPUT_DIR", "outputs/nature_comm_100")
STAGE_DIR = os.path.join(OUTPUT_DIR, "stage_outputs")
REPORT_PATH = os.path.join(OUTPUT_DIR, "全流程汇报文档.docx")

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    hs.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    hs.paragraph_format.space_after = Pt(6)

# ── 辅助函数 ──────────────────────────────────────────────────────────
def load_json(filename):
    """安全加载 JSON 文件."""
    path = os.path.join(STAGE_DIR, filename) if not os.path.isabs(filename) else filename
    if not os.path.exists(path):
        # 尝试 OUTPUT_DIR
        path2 = os.path.join(OUTPUT_DIR, filename)
        if os.path.exists(path2):
            path = path2
        else:
            print(f"  [WARN] 文件不存在: {path}")
            return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def add_code_block(text: str):
    """添加代码块（灰底等宽字体）."""
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        run.element.rPr.append(shd)


def add_json_block(obj, indent=2):
    """添加 JSON 代码块."""
    if isinstance(obj, str):
        add_code_block(obj)
    else:
        text = json.dumps(obj, ensure_ascii=False, indent=indent)
        # 截断过长的 JSON
        lines = text.split("\n")
        if len(lines) > 60:
            lines = lines[:55] + ["  ...", f"  // (共 {len(lines)} 行，已截断)"]
        add_code_block("\n".join(lines))


def add_table(headers, rows, col_widths=None):
    """添加表格."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def add_bullet(text, level=0):
    """添加项目符号."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)


def add_bold_text(bold_part, normal_part):
    """添加加粗前缀 + 正常文本."""
    p = doc.add_paragraph()
    run_b = p.add_run(bold_part)
    run_b.bold = True
    p.add_run(normal_part)


def truncate(s, maxlen=200):
    """截断字符串."""
    if not s:
        return ""
    s = str(s)
    return s[:maxlen] + "..." if len(s) > maxlen else s


# ══════════════════════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("跨学科知识抽取与假设生成系统")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Nature Communications 100篇全流程测试报告")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}")
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  目录概览
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("目录", level=1)
toc_items = [
    "一、项目概述与流程总览",
    "二、Stage 1：学科分类与跨学科筛选",
    "三、Stage 2a：概念抽取",
    "四、Stage 2b：跨学科关系抽取",
    "五、Stage 2c：三级查询生成",
    "六、Stage 2d：三级假设生成",
    "七、Stage 2e：知识图谱构建与指标",
    "八、Stage 3：Benchmark GT 数据集构建",
    "九、Stage 4：P1-P5 五级 Prompt 消融实验",
    "十、Stage 5：多维度评测结果",
    "十一、Stage 6：KG-based 结构化评测",
    "十二、总结与展望",
]
for item in toc_items:
    add_bullet(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  一、项目概述
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("一、项目概述与流程总览", level=1)

doc.add_paragraph(
    "本报告展示了跨学科知识抽取与假设生成系统在 Nature Communications 100 篇论文上的"
    "全流程测试结果。系统从学术论文中自动抽取跨学科概念、关系，生成结构化科学假设，"
    "并通过多维度指标进行评估。"
)

doc.add_heading("1.1 测试数据", level=2)
add_bullet("数据来源: OpenAlex 2025 年 Nature Communications 论文")
add_bullet("筛选条件: Nature 系列期刊 + 有摘要 + 英文")
add_bullet("测试规模: 100 篇论文")

doc.add_heading("1.2 端到端流程", level=2)
add_code_block("""输入 (100篇论文 JSONL)
    ↓
Stage 1: 学科分类 + 跨学科筛选 (~300次 LLM 调用)
    ↓
Stage 2: 三阶段知识抽取 (概念→关系→查询→假设, ~7次/篇)
    ↓
Stage 3: 构建 Benchmark GT 数据集
    ↓
Stage 4: P1-P5 五级 Prompt 假设生成 (消融实验)
    ↓
Stage 5: 多维度评测 (LLM-as-Judge + 文本相似度 + 结构化指标)
    ↓
Stage 6: KG-based 结构化评测""")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  二、Stage 1: 学科分类
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("二、Stage 1：学科分类与跨学科筛选", level=1)

doc.add_paragraph(
    "对 100 篇论文进行层级学科分类，识别主学科与辅学科，筛选出跨学科论文。"
    "分类基于 MSC 层级分类体系，每篇论文约需 3 次 LLM 调用。"
)

stage1_data = load_json("stage1_classification_samples.json")
if stage1_data:
    doc.add_heading("2.1 分类结果样例（3篇）", level=2)
    for i, sample in enumerate(stage1_data[:3]):
        doc.add_heading(f"样例 {i+1}", level=3)
        add_bold_text("标题: ", truncate(sample.get("title", ""), 120))
        add_bold_text("主学科: ", sample.get("primary", ""))
        sec_list = sample.get("secondary_list", sample.get("secondary", []))
        if isinstance(sec_list, list):
            add_bold_text("辅学科: ", ", ".join(sec_list))
        else:
            add_bold_text("辅学科: ", str(sec_list))
        add_bold_text("是否跨学科: ", "是")
        doc.add_paragraph()
else:
    doc.add_paragraph("[分类样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  三、Stage 2a: 概念抽取
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("三、Stage 2a：概念抽取", level=1)

doc.add_paragraph(
    "从论文标题和摘要中抽取专业术语，按主学科和辅学科分组。"
    "每个术语包含标准化形式、原文证据和置信度。"
)

stage2a_data = load_json("stage2a_concepts_samples.json")
if stage2a_data:
    doc.add_heading("3.1 概念抽取样例（3篇）", level=2)
    for i, sample in enumerate(stage2a_data[:3]):
        doc.add_heading(f"样例 {i+1}: {truncate(sample.get('title', ''), 80)}", level=3)
        add_bold_text("主学科: ", sample.get("primary", ""))
        # 主学科概念
        main_concepts = sample.get("概念_主学科", [])
        if main_concepts:
            add_bold_text("主学科概念: ", "")
            rows = []
            for c in main_concepts[:6]:
                if isinstance(c, dict):
                    rows.append([
                        c.get("term", ""),
                        c.get("normalized", ""),
                        truncate(c.get("evidence", ""), 50),
                        str(c.get("confidence", "")),
                    ])
            if rows:
                add_table(["术语", "标准化", "证据", "置信度"], rows, col_widths=[3, 3.5, 6, 1.5])
        # 辅学科概念
        sec_concepts = sample.get("概念_辅学科", {})
        if sec_concepts:
            for disc, concepts in sec_concepts.items():
                add_bold_text(f"辅学科 [{disc}] 概念: ", "")
                rows = []
                for c in (concepts or [])[:4]:
                    if isinstance(c, dict):
                        rows.append([
                            c.get("term", ""),
                            c.get("normalized", ""),
                            str(c.get("confidence", "")),
                        ])
                if rows:
                    add_table(["术语", "标准化", "置信度"], rows, col_widths=[4, 5, 2])
        doc.add_paragraph()
else:
    doc.add_paragraph("[概念抽取样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  四、Stage 2b: 跨学科关系抽取
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("四、Stage 2b：跨学科关系抽取", level=1)

doc.add_paragraph(
    "基于抽取的概念，识别跨学科概念间的语义关系。"
    "系统支持 11 种标准关系类型（如 method_applied_to, improves_metric 等）。"
)

stage2b_data = load_json("stage2b_relations_samples.json")
if stage2b_data:
    doc.add_heading("4.1 关系抽取样例（3篇）", level=2)
    for i, sample in enumerate(stage2b_data[:3]):
        doc.add_heading(f"样例 {i+1}: {truncate(sample.get('title', ''), 80)}", level=3)
        relations = sample.get("跨学科关系", [])
        if relations:
            rows = []
            for r in relations[:5]:
                if isinstance(r, dict):
                    rows.append([
                        r.get("head", ""),
                        r.get("relation", r.get("relation_type", "")),
                        r.get("tail", ""),
                        str(r.get("confidence", "")),
                    ])
            if rows:
                add_table(["头实体", "关系", "尾实体", "置信度"], rows, col_widths=[3.5, 4, 3.5, 1.5])
        # 按辅助学科分类
        by_disc = sample.get("按辅助学科分类", {})
        if by_disc:
            for disc, info in by_disc.items():
                if isinstance(info, dict):
                    rationale = info.get("rationale", "")
                    if rationale:
                        add_bold_text(f"[{disc}] ", truncate(rationale, 150))
        doc.add_paragraph()
else:
    doc.add_paragraph("[关系抽取样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  五、Stage 2c: 查询生成
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("五、Stage 2c：三级查询生成", level=1)

doc.add_paragraph(
    "基于抽取的概念和关系，生成三级层次化查询："
    "一级（宏观核心问题）→ 二级（按辅学科维度）→ 三级（细粒度操作问题）。"
)

stage2c_data = load_json("stage2c_query_samples.json")
if stage2c_data:
    doc.add_heading("5.1 查询生成样例（3篇）", level=2)
    for i, sample in enumerate(stage2c_data[:3]):
        doc.add_heading(f"样例 {i+1}: {truncate(sample.get('title', ''), 80)}", level=3)
        queries = sample.get("查询", {})
        if queries:
            l1 = queries.get("一级", "")
            if l1:
                add_bold_text("一级查询: ", truncate(l1, 200))
            l2 = queries.get("二级", [])
            if l2:
                add_bold_text("二级查询: ", "")
                for q in l2[:3]:
                    add_bullet(truncate(str(q), 150))
            l3 = queries.get("三级", [])
            if l3:
                add_bold_text("三级查询: ", "")
                for q in l3[:3]:
                    add_bullet(truncate(str(q), 150))
        doc.add_paragraph()
else:
    doc.add_paragraph("[查询生成样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  六、Stage 2d: 假设生成
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("六、Stage 2d：三级假设生成", level=1)

doc.add_paragraph(
    "基于概念、关系和查询，生成三级结构化假设。"
    "每条假设为 3 步链式推理路径（head → relation → tail），附带总结。"
)

stage2d_data = load_json("stage2d_hypothesis_samples.json")
if stage2d_data:
    doc.add_heading("6.1 假设生成样例（3篇）", level=2)
    for i, sample in enumerate(stage2d_data[:3]):
        doc.add_heading(f"样例 {i+1}: {truncate(sample.get('title', ''), 80)}", level=3)
        for level_key, sum_key in [("假设_一级", "假设_一级总结"),
                                    ("假设_二级", "假设_二级总结"),
                                    ("假设_三级", "假设_三级总结")]:
            paths = sample.get(level_key, [])
            summaries = sample.get(sum_key, [])
            level_label = level_key.replace("假设_", "")
            if paths:
                add_bold_text(f"{level_label}假设 ({len(paths)} 条路径): ", "")
                # 展示第一条路径
                path = paths[0] if paths else []
                if isinstance(path, list):
                    rows = []
                    for step in path[:3]:
                        if isinstance(step, dict):
                            rows.append([
                                str(step.get("step", "")),
                                step.get("head", step.get("头实体", "")),
                                step.get("relation", step.get("关系", "")),
                                step.get("tail", step.get("尾实体", "")),
                            ])
                    if rows:
                        add_table(["步骤", "头实体", "关系", "尾实体"], rows, col_widths=[1.5, 4, 4, 4])
            if summaries:
                add_bold_text(f"{level_label}总结: ", "")
                for s in summaries[:2]:
                    add_bullet(truncate(str(s), 200))
        doc.add_paragraph()
else:
    doc.add_paragraph("[假设生成样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  七、Stage 2e: 图谱指标
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("七、Stage 2e：知识图谱构建与指标", level=1)

doc.add_paragraph(
    "基于抽取的概念和关系构建知识图谱，计算 15+ 维度的图谱质量指标。"
)

stage2e_data = load_json("stage2e_graph_metrics_samples.json")
if stage2e_data:
    doc.add_heading("7.1 图谱指标样例（3篇）", level=2)
    for i, sample in enumerate(stage2e_data[:3]):
        doc.add_heading(f"样例 {i+1}: {truncate(sample.get('title', ''), 80)}", level=3)
        graph = sample.get("graph", {})
        metrics = sample.get("metrics", {})
        nodes = graph.get("nodes", [])
        edges = graph.get("edges", [])
        add_bold_text("图谱规模: ", f"{len(nodes)} 个节点, {len(edges)} 条边")
        if metrics:
            rows = []
            for k, v in list(metrics.items())[:8]:
                rows.append([k, f"{v:.4f}" if isinstance(v, float) else str(v)])
            if rows:
                add_table(["指标", "值"], rows, col_widths=[6, 4])
        doc.add_paragraph()
else:
    doc.add_paragraph("[图谱指标样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  八、Stage 3: GT 构建
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("八、Stage 3：Benchmark GT 数据集构建", level=1)

doc.add_paragraph(
    "从抽取结果中构建 Evidence-Grounded 的 Ground Truth 数据集，"
    "用于后续评测。GT 包含术语、关系和跨学科路径。"
)

stage3_data = load_json("stage3_gt_benchmark_samples.json")
if stage3_data:
    doc.add_heading("8.1 GT 样例（3条）", level=2)
    for i, sample in enumerate(stage3_data[:3]):
        doc.add_heading(f"样例 {i+1}", level=3)
        inp = sample.get("input", {})
        gt = sample.get("ground_truth", {})
        add_bold_text("标题: ", truncate(inp.get("title", sample.get("id", "")), 120))
        add_bold_text("主学科: ", inp.get("primary_discipline", ""))
        sec = inp.get("secondary_disciplines", [])
        add_bold_text("辅学科: ", ", ".join(sec) if isinstance(sec, list) else str(sec))
        terms = gt.get("terms", [])
        add_bold_text(f"GT 术语数: ", str(len(terms)))
        rels = gt.get("relations", [])
        add_bold_text(f"GT 关系数: ", str(len(rels)))
        paths = gt.get("paths", [])
        add_bold_text(f"GT 路径数: ", str(len(paths)))
        doc.add_paragraph()
else:
    doc.add_paragraph("[GT 样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  九、Stage 4: P1-P5 消融实验
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("九、Stage 4：P1-P5 五级 Prompt 消融实验", level=1)

doc.add_paragraph(
    "对测试集论文，分别使用 P1-P5 五种不同信息量的 Prompt 生成假设，"
    "用于评估不同引导程度对跨学科假设生成质量的影响。"
)

doc.add_heading("9.1 P1-P5 级别说明", level=2)
add_table(
    ["级别", "Query", "论文信息", "结构化知识", "格式"],
    [
        ["P1", "L1 only", "无", "无", "自由文本"],
        ["P2", "L1 + L2", "abstract + 学科角色", "无", "自由文本"],
        ["P3", "L1+L2+L3", "abstract + 学科角色", "概念列表", "半结构化"],
        ["P4", "L1+L2+L3", "abstract + 学科角色", "概念 + 关系", "半结构化"],
        ["P5", "L1+L2+L3", "完整结构化摘要", "概念+关系+路径", "3-step 路径"],
    ],
    col_widths=[1.5, 2.5, 3.5, 3.5, 3],
)

stage4_data = load_json("stage4_p1p5_samples.json")
if stage4_data:
    doc.add_heading("9.2 P1-P5 生成样例", level=2)
    for level_name in ["P1", "P2", "P3", "P4", "P5"]:
        samples = stage4_data.get(level_name, [])
        if samples:
            doc.add_heading(f"{level_name} 样例", level=3)
            sample = samples[0]
            add_bold_text("论文: ", truncate(sample.get("title", ""), 100))
            add_bold_text("主学科: ", sample.get("primary", ""))
            add_bold_text(f"耗时: ", f"{sample.get('elapsed', 0)} 秒")
            hyp_text = sample.get("hypothesis_text", "")
            if hyp_text:
                add_bold_text("生成假设: ", "")
                add_code_block(truncate(hyp_text, 800))
            doc.add_paragraph()
else:
    doc.add_paragraph("[P1-P5 样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  十、Stage 5: 多维度评测
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("十、Stage 5：多维度评测结果", level=1)

doc.add_paragraph(
    "对 P1-P5 各级假设进行多维度评测，包括 LLM-as-Judge 主观评分、"
    "文本相似度和结构化指标。"
)

stage5_data = load_json("stage5_eval_samples.json")
if stage5_data:
    # 聚合汇总
    aggregated = stage5_data.get("aggregated_summary", {})
    if aggregated:
        doc.add_heading("10.1 P1-P5 方法对比汇总", level=2)
        headers = ["方法"]
        first_method = next(iter(aggregated.values()), {})
        metric_keys = list(first_method.keys())[:8]
        headers.extend(metric_keys)
        rows = []
        for method_name in ["P1", "P2", "P3", "P4", "P5"]:
            method_data = aggregated.get(method_name, {})
            if method_data:
                row = [method_name]
                for mk in metric_keys:
                    val = method_data.get(mk, "")
                    if isinstance(val, float):
                        row.append(f"{val:.3f}")
                    else:
                        row.append(str(val))
                rows.append(row)
        if rows:
            add_table(headers, rows)

    # 详细结果样例
    detailed = stage5_data.get("detailed_results", [])
    if detailed:
        doc.add_heading("10.2 评测详细结果样例", level=2)
        for i, result in enumerate(detailed[:3]):
            doc.add_heading(f"评测样例 {i+1}", level=3)
            add_bold_text("方法: ", result.get("method_name", ""))
            add_bold_text("论文ID: ", result.get("paper_id", ""))
            scores = result.get("scores", result.get("metrics", {}))
            if scores:
                rows = []
                for k, v in list(scores.items())[:10]:
                    rows.append([k, f"{v:.4f}" if isinstance(v, float) else str(v)])
                if rows:
                    add_table(["评测维度", "分值"], rows, col_widths=[6, 4])
            doc.add_paragraph()
else:
    doc.add_paragraph("[评测样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  十一、Stage 6: KG-based 评测
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("十一、Stage 6：KG-based 结构化评测", level=1)

doc.add_paragraph(
    "针对 P5 结构化路径假设，使用 GT 知识图谱进行深度评测，"
    "评估假设路径与 GT 的术语覆盖、关系对齐和路径匹配。"
)

stage6_data = load_json("stage6_kg_eval_samples.json")
if stage6_data and isinstance(stage6_data, list):
    doc.add_heading("11.1 KG 评测样例", level=2)
    for i, sample in enumerate(stage6_data[:3]):
        doc.add_heading(f"样例 {i+1}", level=3)
        if isinstance(sample, dict):
            scores = sample.get("scores", {})
            if scores:
                rows = []
                for k, v in list(scores.items())[:10]:
                    rows.append([k, f"{v:.4f}" if isinstance(v, float) else str(v)])
                if rows:
                    add_table(["评测维度", "分值"], rows, col_widths=[6, 4])
            else:
                add_json_block(sample)
        doc.add_paragraph()
elif stage6_data:
    add_json_block(stage6_data)
else:
    doc.add_paragraph("[KG 评测样例文件未找到]")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  十二、总结
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("十二、总结与展望", level=1)

doc.add_paragraph(
    "本次测试在 Nature Communications 100 篇论文上完整运行了跨学科知识抽取与假设生成系统的全流程。"
    "测试覆盖了从学科分类、概念/关系抽取、查询/假设生成，到 P1-P5 消融实验和多维度评测的所有阶段。"
)

doc.add_heading("12.1 主要发现", level=2)
add_bullet("系统能够有效识别跨学科论文并抽取结构化知识")
add_bullet("三级假设生成（L1/L2/L3）覆盖了从宏观到细粒度的不同层次")
add_bullet("P1-P5 消融实验验证了结构化信息对假设生成质量的提升作用")
add_bullet("多维度评测框架提供了全面的假设质量评估")

doc.add_heading("12.2 后续方向", level=2)
add_bullet("扩大测试规模，覆盖更多学科领域")
add_bullet("优化 Prompt 设计，提升假设的创新性和可行性")
add_bullet("引入人工评估，与自动评测指标进行对比验证")
add_bullet("探索更多 baseline 方法（如 IdeaBench、SciMON 等）的对比")

# ── 保存 ──
doc.save(REPORT_PATH)
print(f"汇报文档已保存至: {REPORT_PATH}")
